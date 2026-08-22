from __future__ import annotations

import io
from datetime import datetime, timedelta, timezone
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Brasília é UTC-3 fixo (o Brasil não tem horário de verão desde 2019), então um
# offset fixo é exato e evita depender do banco de fusos (tzdata) no Windows.
FUSO_BRASILIA = timezone(timedelta(hours=-3))

_COR_CLASSIF = {
    'ALTO': RGBColor(0xC0, 0x39, 0x2B),
    'MEDIO': RGBColor(0xE6, 0x7E, 0x22),
    'BAIXO': RGBColor(0x27, 0xAE, 0x60),
}

_ORIGEM_TEXTO = {
    'llm': 'Análise escrita por IA (Google Gemini) sobre scores determinísticos.',
    'prompt_apenas': 'Análise por template (sem chave de IA configurada).',
    'fallback': 'Análise por template (provedor de IA indisponível no momento).',
}


def para_brasilia(iso_utc: str | None) -> str:
    """Converte um timestamp ISO em UTC para 'DD/MM/AAAA HH:MM' no horário de Brasília."""
    if not iso_utc:
        return '—'
    try:
        dt = datetime.fromisoformat(str(iso_utc).replace('Z', '+00:00'))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(FUSO_BRASILIA).strftime('%d/%m/%Y %H:%M')
    except (ValueError, TypeError):
        return str(iso_utc)


def _add_par(doc, texto: str, negrito: bool = False, tamanho: int | None = None):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrito
    if tamanho:
        run.font.size = Pt(tamanho)
    return p


def _add_eixo(doc, titulo: str, score: Any, classif: str, justificativa: str) -> None:
    doc.add_heading(titulo, level=2)

    p = doc.add_paragraph()
    p.add_run('Score: ').bold = True
    p.add_run(f'{score}  ')
    tag = p.add_run(f'({classif})')
    tag.bold = True
    tag.font.color.rgb = _COR_CLASSIF.get(str(classif).upper(), RGBColor(0x33, 0x33, 0x33))

    if justificativa:
        doc.add_paragraph(justificativa)


def montar_docx(relatorio: Dict[str, Any], eventos: List[Dict[str, Any]]) -> bytes:
    """Monta o relatório de risco como um documento Word (.docx) e devolve os bytes."""
    doc = Document()

    doc.add_heading('Relatório de Risco', level=0)

    # Cabeçalho de identificação
    gerado = datetime.now(FUSO_BRASILIA).strftime('%d/%m/%Y %H:%M')
    sub = doc.add_paragraph()
    sub.add_run(f"Dispositivo: {relatorio.get('dispositivo', '—')}").bold = True
    sub.add_run(
        f"   ·   Período: {relatorio.get('periodo_dias', '—')} dia(s)"
        f"   ·   Gerado em: {gerado} (Brasília)"
    )

    origem = relatorio.get('origem_da_analise', '')
    nota = doc.add_paragraph()
    r = nota.add_run(_ORIGEM_TEXTO.get(origem, f'Origem da análise: {origem}'))
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Scores
    _add_eixo(
        doc, '🔓 Furto',
        relatorio.get('score_furto', '—'),
        relatorio.get('classificacao_furto', '—'),
        relatorio.get('justificativa_furto', ''),
    )
    _add_eixo(
        doc, '🔥 Incêndio',
        relatorio.get('score_incendio', '—'),
        relatorio.get('classificacao_incendio', '—'),
        relatorio.get('justificativa_incendio', ''),
    )

    # Recomendações
    recomendacoes = relatorio.get('recomendacoes') or []
    if recomendacoes:
        doc.add_heading('Recomendações', level=2)
        for item in recomendacoes:
            doc.add_paragraph(str(item), style='List Bullet')

    # Limitações
    limitacoes = relatorio.get('limitacoes')
    if limitacoes:
        doc.add_heading('Limitações', level=2)
        doc.add_paragraph(str(limitacoes))

    # Eventos do período (horários já em Brasília)
    doc.add_heading('Eventos no período', level=2)
    if eventos:
        tabela = doc.add_table(rows=1, cols=3)
        tabela.style = 'Light Grid Accent 1'
        cab = tabela.rows[0].cells
        cab[0].paragraphs[0].add_run('Data/hora (Brasília)').bold = True
        cab[1].paragraphs[0].add_run('Tipo').bold = True
        cab[2].paragraphs[0].add_run('Severidade').bold = True
        for ev in eventos:
            linha = tabela.add_row().cells
            linha[0].text = para_brasilia(ev.get('criado_em'))
            linha[1].text = str(ev.get('tipo', '—'))
            linha[2].text = str(ev.get('severidade', '—'))
    else:
        doc.add_paragraph('Nenhum evento registrado no período.')

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
